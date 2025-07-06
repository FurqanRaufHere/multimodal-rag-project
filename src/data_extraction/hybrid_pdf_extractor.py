import fitz  # PyMuPDF
import easyocr
import cv2
import numpy as np
from pdf2image import convert_from_path
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from typing import List, Dict, Union
import json
import os
import io
import base64
from PIL import Image, ImageDraw
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle
import pytesseract
from transformers import BlipProcessor, BlipForConditionalGeneration
import torch

class EnhancedFileExtractor:
    def __init__(self):
        self.reader = None
        self.blip_processor = None
        self.blip_model = None
        os.makedirs("output/images", exist_ok=True)
        os.makedirs("output/pages", exist_ok=True)

    def _init_easyocr(self):
        if self.reader is None:
            self.reader = easyocr.Reader(['en'])

    def _init_blip(self):
        """Initialize BLIP model for image captioning"""
        if self.blip_processor is None:
            self.blip_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
            self.blip_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")

    def _detect_charts_in_image(self, image_path: str) -> Dict:
        """Detect if image contains charts/diagrams using computer vision"""
        img = cv2.imread(image_path)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Detect lines (common in charts)
        edges = cv2.Canny(gray, 50, 150, apertureSize=3)
        lines = cv2.HoughLines(edges, 1, np.pi/180, threshold=100)
        
        # Detect rectangles (bars in bar charts)
        contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        rectangles = []
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            if w > 20 and h > 20:  # Filter small noise
                rectangles.append((x, y, w, h))
        
        # Simple heuristics for chart detection
        has_lines = lines is not None and len(lines) > 5
        has_rectangles = len(rectangles) > 3
        
        chart_type = "unknown"
        if has_rectangles and has_lines:
            chart_type = "bar_chart"
        elif has_lines:
            chart_type = "line_chart"
        elif has_rectangles:
            chart_type = "diagram"
            
        return {
            "is_chart": has_lines or has_rectangles,
            "chart_type": chart_type,
            "line_count": len(lines) if lines is not None else 0,
            "rectangle_count": len(rectangles)
        }

    def _extract_page_as_image(self, page, page_num: int, pdf_path: str) -> str:
        """Convert PDF page to image to capture vector graphics"""
        # Get page as image
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
        img_data = pix.tobytes("png")
        
        # Save full page image
        page_img_path = f"output/pages/{os.path.basename(pdf_path)}_page_{page_num + 1}.png"
        with open(page_img_path, "wb") as f:
            f.write(img_data)
        
        return page_img_path

    def _generate_image_description(self, image_path: str) -> str:
        """Generate description of image using BLIP model"""
        try:
            self._init_blip()
            image = Image.open(image_path).convert('RGB')
            
            # Generate caption
            inputs = self.blip_processor(image, return_tensors="pt")
            out = self.blip_model.generate(**inputs, max_length=50)
            caption = self.blip_processor.decode(out[0], skip_special_tokens=True)
            
            return caption
        except Exception as e:
            print(f"BLIP captioning failed: {str(e)}")
            return ""

    def _extract_text_from_regions(self, image_path: str) -> List[Dict]:
        """Extract text from different regions of the image"""
        self._init_easyocr()
        img = cv2.imread(image_path)
        
        # Get OCR results with bounding boxes
        ocr_results = self.reader.readtext(img, paragraph=False)
        
        text_regions = []
        for bbox, text, confidence in ocr_results:
            if confidence > 0.5:  # Filter low confidence
                text_regions.append({
                    "text": text,
                    "confidence": confidence,
                    "bbox": bbox
                })
        
        return text_regions

    # ==================== ENHANCED PDF PROCESSING ====================
    def extract_pdf_content(self, pdf_path: str) -> List[Dict]:
        """Enhanced PDF extraction with better chart/diagram detection"""
        doc = fitz.open(pdf_path)
        chunks = []
        
        for page_num, page in enumerate(doc):
            print(f"Processing page {page_num + 1}/{len(doc)}")
            
            # 1. Extract text content
            text = page.get_text("text")
            if text.strip():
                chunks.append({
                    "content": text,
                    "type": "text",
                    "page": page_num + 1,
                    "source": pdf_path,
                    "metadata": {"word_count": len(text.split())}
                })
            
            # 2. Extract tables
            tables = page.find_tables()
            if tables:
                for table_idx, table in enumerate(tables):
                    table_data = table.extract()
                    chunks.append({
                        "content": str(table_data),
                        "type": "table",
                        "page": page_num + 1,
                        "source": pdf_path,
                        "metadata": {
                            "table_id": table_idx,
                            "rows": len(table_data),
                            "cols": len(table_data[0]) if table_data else 0
                        }
                    })
            
            # 3. Extract embedded images
            for img_index, img in enumerate(page.get_images(full=True)):
                try:
                    base_image = doc.extract_image(img[0])
                    img_data = base_image["image"]
                    
                    # Save image
                    img_path = f"output/images/{os.path.basename(pdf_path)}_p{page_num+1}_embedded_{img_index}.png"
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                    
                    # Analyze image
                    chart_info = self._detect_charts_in_image(img_path)
                    text_regions = self._extract_text_from_regions(img_path)
                    description = self._generate_image_description(img_path)
                    
                    chunks.append({
                        "image_path": img_path,
                        "content": " ".join([region["text"] for region in text_regions]),
                        "type": "embedded_image",
                        "page": page_num + 1,
                        "source": pdf_path,
                        "metadata": {
                            "dimensions": f"{base_image['width']}x{base_image['height']}",
                            "is_chart": chart_info["is_chart"],
                            "chart_type": chart_info["chart_type"],
                            "description": description,
                            "text_regions": text_regions
                        }
                    })
                except Exception as e:
                    print(f"Embedded image processing error (Page {page_num+1}): {str(e)}")
            
            # 4. Extract full page as image (to capture vector graphics)
            page_img_path = self._extract_page_as_image(page, page_num, pdf_path)
            
            # Analyze full page for charts/diagrams
            chart_info = self._detect_charts_in_image(page_img_path)
            
            # Only add if page seems to contain significant visual content
            if chart_info["is_chart"] or chart_info["line_count"] > 10:
                text_regions = self._extract_text_from_regions(page_img_path)
                description = self._generate_image_description(page_img_path)
                
                chunks.append({
                    "image_path": page_img_path,
                    "content": " ".join([region["text"] for region in text_regions]),
                    "type": "page_image",
                    "page": page_num + 1,
                    "source": pdf_path,
                    "metadata": {
                        "is_chart": chart_info["is_chart"],
                        "chart_type": chart_info["chart_type"],
                        "description": description,
                        "visual_elements": {
                            "lines": chart_info["line_count"],
                            "rectangles": chart_info["rectangle_count"]
                        }
                    }
                })
        
        doc.close()
        return chunks

    # ==================== ENHANCED DOCX PROCESSING ====================
    def extract_docx_content(self, docx_path: str) -> List[Dict]:
        """Enhanced DOCX extraction"""
        doc = Document(docx_path)
        chunks = []
        
        # 1. Extract paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                chunks.append({
                    "content": para.text,
                    "type": "text",
                    "source": docx_path,
                    "metadata": {
                        "paragraph_id": para_idx,
                        "style": para.style.name if para.style else None
                    }
                })
        
        # 2. Extract tables
        for table_num, table in enumerate(doc.tables):
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            chunks.append({
                "content": str(table_data),
                "type": "table",
                "source": docx_path,
                "metadata": {
                    "table_id": table_num,
                    "rows": len(table_data),
                    "cols": len(table_data[0]) if table_data else 0
                }
            })
        
        # 3. Extract images with enhanced processing
        try:
            for rel_idx, rel in enumerate(doc.part.rels.values()):
                if "image" in rel.target_ref:
                    img_data = rel.target_part.blob
                    img_path = f"output/images/{os.path.basename(docx_path)}_img_{rel_idx}.png"
                    
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                    
                    # Analyze image
                    chart_info = self._detect_charts_in_image(img_path)
                    text_regions = self._extract_text_from_regions(img_path)
                    description = self._generate_image_description(img_path)
                    
                    chunks.append({
                        "image_path": img_path,
                        "content": " ".join([region["text"] for region in text_regions]),
                        "type": "docx_image",
                        "source": docx_path,
                        "metadata": {
                            "is_chart": chart_info["is_chart"],
                            "chart_type": chart_info["chart_type"],
                            "description": description,
                            "text_regions": text_regions
                        }
                    })
        except Exception as e:
            print(f"DOCX image extraction error: {str(e)}")
        
        return chunks

    # ==================== MAIN PROCESSOR ===================
    def process_files(self, file_paths: List[str]) -> List[Dict]:
        """Process all files with enhanced multimodal extraction"""
        all_data = []
        
        for file_path in file_paths:
            print(f"\nProcessing: {file_path}")
            try:
                if file_path.lower().endswith('.pdf'):
                    chunks = self.extract_pdf_content(file_path)
                elif file_path.lower().endswith('.docx'):
                    chunks = self.extract_docx_content(file_path)
                else:
                    print(f"Unsupported file type: {file_path}")
                    continue
                
                all_data.extend(chunks)
                print(f"Extracted {len(chunks)} chunks from {file_path}")
                
            except Exception as e:
                print(f"Failed to process {file_path}: {str(e)}")
        
        return all_data

    def get_statistics(self, data: List[Dict]) -> Dict:
        """Get statistics about extracted data"""
        stats = {
            "total_chunks": len(data),
            "by_type": {},
            "by_source": {},
            "charts_detected": 0,
            "images_with_text": 0
        }
        
        for item in data:
            # Count by type
            item_type = item.get("type", "unknown")
            stats["by_type"][item_type] = stats["by_type"].get(item_type, 0) + 1
            
            # Count by source
            source = os.path.basename(item.get("source", "unknown"))
            stats["by_source"][source] = stats["by_source"].get(source, 0) + 1
            
            # Count charts
            if item.get("metadata", {}).get("is_chart"):
                stats["charts_detected"] += 1
            
            # Count images with text
            if item.get("type") in ["embedded_image", "page_image", "docx_image"] and item.get("content"):
                stats["images_with_text"] += 1
        
        return stats

# Usage
import numpy as np

def convert_np_int_to_int(obj):
    if isinstance(obj, list):
        return [convert_np_int_to_int(i) for i in obj]
    elif isinstance(obj, dict):
        return {k: convert_np_int_to_int(v) for k, v in obj.items()}
    elif isinstance(obj, np.integer):
        return int(obj)
    else:
        return obj

if __name__ == "__main__":
    extractor = EnhancedFileExtractor()
    files = [
        "data/Enunciado-VC-TP-2025-EN.pdf",
        "data/Milestones_Breakdown.pdf", 
        "data/PDC_Report.docx",
        "data/SDA_Project_Document.docx"
    ]
    
    print("Starting enhanced multimodal extraction...")
    result = extractor.process_files(files)
    
    # Convert numpy int types to int for JSON serialization
    result = convert_np_int_to_int(result)
    
    # Save results
    with open("output/extracted_data.json", "w", encoding='utf-8') as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    
    # Print statistics
    stats = extractor.get_statistics(result)
    print("\n" + "="*50)
    print("EXTRACTION STATISTICS")
    print("="*50)
    print(f"Total chunks extracted: {stats['total_chunks']}")
    print(f"Charts detected: {stats['charts_detected']}")
    print(f"Images with text: {stats['images_with_text']}")
    
    print("\nBy Type:")
    for type_name, count in stats['by_type'].items():
        print(f"  {type_name}: {count}")
    
    print("\nBy Source:")
    for source, count in stats['by_source'].items():
        print(f"  {source}: {count}")
    
    print(f"\nImages saved to: output/images/")
    print(f"Page images saved to: output/pages/")
    print(f"Data saved to: output/extracted_data.json")
